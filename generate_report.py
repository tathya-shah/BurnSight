import sys
import subprocess

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    install('python-docx')
    import docx

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('Hacka-Mined: Complete Features & Analysis Report', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_heading('1. Problem Statement', level=1)
doc.add_paragraph(
    "Startups and small businesses often fail due to poor cash flow management, lack of visibility into \"zombie spend\" (recurring expenses that provide no value), "
    "and a critical misunderstanding of their true financial runway. Traditional financial dashboards simply categorize income and expenses, "
    "leaving founders to interpret the data and make complex financial projections themselves. There is a strong need for an intelligent, "
    "founder-centric tool that acts as an automated CFO—analyzing raw bank transactions, projecting survival timelines, and offering highly actionable advice to prevent failure."
)

doc.add_heading('2. What Unique Things Have We Done?', level=1)
p = doc.add_paragraph()
p.add_run('1. Actionable AI CFO Insights: ').bold = True
p.add_run('Instead of just showing charts, the system translates raw transaction data into clear English recommendations. It actively identifies "Zombie Spend" and predicts the exact financial impact (e.g., "Reducing AWS costs by 10% extends runway by 1.2 months").\n')

p.add_run('2. The Survival Simulator: ').bold = True
p.add_run('A highly interactive scenario-modeling engine. Founders can adjust sliders for "Hiring Employees", "Marketing Budget", "Revenue Growth", and "Fundraise Amount" to instantly see how these decisions impact their Runway and Risk Score in real-time.\n')

p.add_run('3. Founder-Centric Metrics: ').bold = True
p.add_run('The platform tracks startup-specific survival metrics like "Burn Velocity", "Net Burn", and proprietary "Risk Scores" (0-100), rather than generic accounting terms.\n')

p.add_run('4. Resilient Hybrid Engine: ').bold = True
p.add_run('We engineered a dual-layer analysis system. If the AI service is unavailable, it seamlessly falls back to a deterministic algorithm to ensure the founder always has access to their critical runway calculations and basic categorization without interruption.')

doc.add_heading('3. AI Integration Detail', level=1)
doc.add_paragraph('We integrated Groq\'s high-speed inference API utilizing the LLaMA3-8B-8192 model to power the core intelligence engine. The AI operates under a strict "Startup CFO" system prompt.')
p2 = doc.add_paragraph()
p2.add_run('Key AI Capabilities:\n').bold = True
p2.add_run('- ')
p2.add_run('Unstructured Data Parsing: ').bold = True
p2.add_run('Processes raw bank and Stripe transaction descriptions to accurately categorize spend.\n')
p2.add_run('- ')
p2.add_run('Semantic Contextualization: ').bold = True
p2.add_run('Understands the semantic meaning behind transactions to detect wasteful "Zombie Spend."\n')
p2.add_run('- ')
p2.add_run('Structured JSON Generation: ').bold = True
p2.add_run('The AI is rigorously constrained to output validated JSON containing risk levels, top expense categories, and monthly breakdowns for seamless UI integration.\n')

doc.add_heading('4. Complete Features Pipeline', level=1)
doc.add_paragraph('• CSV Data Ingestion: Users can drag and drop bank statement CSVs along with their current cash balance.')
doc.add_paragraph('• Burn Velocity Charting: Visualizes historical net burn trends month-over-month.')
doc.add_paragraph('• Revenue vs. Expense Tracking: Detailed graphical breakdown of cash flow.')
doc.add_paragraph('• Collapse Warnings: Alerts the founder when runway crosses critical thresholds (e.g., under 6 months).')
doc.add_paragraph('• PDF/UI Exporting: Uses canvas and jsPDF to allow founders to capture and share their financial health reports with investors (via html-to-image).')
doc.add_paragraph('• Modern Tech Stack: Built on Next.js 16 (App Router), React 19, Tailwind CSS v4, Framer Motion for premium animations, and Supabase for backend scalable data.')

doc.save('C:/Users/Lenovo/Downloads/Hacka-Mined-main (1)/Hacka-Mined-main/Features_Report.docx')
print("Successfully generated Features_Report.docx")
